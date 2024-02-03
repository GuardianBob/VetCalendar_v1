

# NOTE: 2021 SChedule format

# from docx import Document
# import re

# wordDoc = Document('SEP21_schedule.docx')
# # wordDoc = Document('Aug 2022(CORRECT).docx')
# month = ''
# year = ''
# shifts = []

# month_list = [
#   "january", "jan",
#   "february", "feb",
#   "march", "mar",
#   "april", "apr",
#   "may",
#   "june", "jun",
#   "july", "jul",
#   "august", "aug",
#   "september", "sept",
#   "october", "oct",
#   "november", "nov",
#   "december", "dec"
# ]

# day_list = [
#   "monday", "mon",
#   "tuesday", "tues",
#   "wednesday", "wed",
#   "thursday", "thurs",
#   "friday", "fri",
#   "saturday", "sat",
#   "sunday", "sun"
# ]

# for table in wordDoc.tables:
#     date = []
#     j = 0
#     for row in table.rows:
#       row_text = ''
#       i = 0
#       k = 0
#       for cell in row.cells:
#         row_text = row_text + cell.text + ","
#         if cell.text.lower() in day_list:
#           j = 0
#           # print('reading cells')
        
#         if j % 3 == 1:
#             date.append(cell.text)
#         else:
#           split_cell = cell.text.split()
#           for x in split_cell:
#             if x.lower() in month_list:
#               month = split_cell[0]
#               year = split_cell[1]
#             if j % 3 == 2:
#                 if x == "MA":
#                   shifts.append(["Day Shift", month + "/" + date[i] + "/" + "20" + year, "6:00 AM"])
#                 if x == "(MA)":
#                   shifts.append(["Swing Shift 1", month + "/" + date[i] + "/" + "20" + year, "10:00 AM"])
#             if j % 3 == 0:
#                 if x == "MA":
#                   shifts.append(["Swing Shift 2", month + "/" + date[i] + "/" + "20" + year, "2:00 PM"])
#                 if x == "(MA)":
#                   shifts.append(["Overnight", month + "/" + date[i] + "/" + "20" + year, "6:00 PM"])

#         i += 1

#       print(row_text, "row", j)
#       # print("date row", date)
#       j += 1
#       i = 0
#       k = 0
#       if j % 3 == 1: 
#         date = []

# print(shifts)

# NOTE: 2022 SChedule format

# ==== Google Calendar API reqs =====
from __future__ import print_function

import datetime, pytz, sys, traceback
from datetime import timedelta
import os.path
from .models import Calendar
from django.utils import timezone
from django.http import JsonResponse


# ================================================================

from docx import Document
# from tkinter import filedialog, simpledialog
import csv, json

TIMEZONE = pytz.timezone('America/Los_Angeles')

month_variables = ["01", "02", "03", "04", "05", "06", "07", "08", "09", "10", "11", "12"]

# wordDoc = Document(filedialog.askopenfilename( filetypes = ( (".docx .doc files", "*.docx *.doc"),("All files", "*.*") ) ))

month_abbrev = ["Jan", "Feb", "Mar", "Apr", "May", "Jun", "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"]

month_list = [
    "january", "jan",
    "february", "feb",
    "march", "mar",
    "april", "apr",
    "may",
    "june", "jun",
    "july", "jul",
    "august", "aug",
    "september", "sept",
    "october", "oct",
    "november", "nov",
    "december", "dec"
  ]

day_list = [
    "monday", "mon",
    "tuesday", "tues",
    "wednesday", "wed",
    "thursday", "thurs",
    "friday", "fri",
    "saturday", "sat",
    "sunday", "sun"
  ]

def trace_error(e, isForm=False):
    exc_type, exc_value, exc_traceback = sys.exc_info()
    filename, line_number, func_name, text = traceback.extract_tb(exc_traceback)[0]
    print(f"An error occurred in file {filename} on line {line_number} in {func_name}(): {text}")
    print("Error: ", e)
    if isForm:
        return JsonResponse({'message':'Form is invalid'}, status=500)
    return JsonResponse({'message':'Something went wrong'}, status=500)

def convert_schedule(schedule, user, month, year):
  user_month = month
  user_year = year
  wordDoc = Document(schedule)
  while not user_month in month_variables:
    # user_month = simpledialog.askstring(title="Month", prompt="Please enter the 2-digit month")
    user_month = "08"

  while not int(user_year) > 2021:
    # user_year = simpledialog.askstring(title="Year", prompt="Please enter the 4-digit year")
    user_year = "2022"
  # wordDoc = Document('SEP21_schedule.docx')
  # wordDoc = Document('Aug 2022(CORRECT).docx')
  month = user_month
  year = "2022"
  shifts = []
  for table in wordDoc.tables:
      date = []
      j = 0
      for row in table.rows:
        row_text = ''
        i = 0
        k = 0
        shift = ''
        time = ''
        for cell in row.cells:
          row_text = row_text + cell.text + ","
          if cell.text.lower() in day_list:
            j = 0
            # print('reading cells')
          
          if j % 6 == 1:
              date.append(cell.text)
          else:
            if cell.text in month_list:
              month = cell.text
            if i % 2 == 0:
              shift = cell.text
              if j % 6 == 2:
                time = "07:00"
              if j % 6 == 3:
                time = "10:00"
              if j % 6 == 4:
                time = "14:00"
              if j % 6 == 5:
                time = "18:00"
            # if "MA" in cell.text:
            # print(user)
            if user in cell.text:
              shift_start = datetime.datetime(
                int(user_year),
                int(user_month),
                int(date[i]),
                int(time[:2]),
                00,
                00
              )
              # print(shift_start)
              shifts.append({
                "shift": shift, 
                "date": user_year + "-" + user_month + "-" + date[i], 
                "time": time, 
                "shift_start": str(shift_start),
                "shift_end": str(shift_start + timedelta(hours=12)),
              })
                
          i += 1

        # print(row_text, "row", j)
        # print("date row", date)
        j += 1
        i = 0
        k = 0
        if j % 6 == 0: 
          date = []

  header = ['Subject', 'Start date', 'Start time']
  # with open(f'schedule_{user_month}-{user_year}.csv', 'w', encoding='UTF8', newline='\n') as f:
  #   writer = csv.writer(f)
  #   writer.writerow(header)
  #   for shift in shifts:
  #     writer.writerow(shift)

  # f.close()
  # start_time = datetime.datetime(2023,6,21,3,45,00)
  # end_time = start_time + timedelta(hours=12)
  # print("start: ", start_time, "end: ", end_time)
  # print(shifts)
  
  events = add_shifts(shifts)
  json_shifts = json.dumps(events)
  return json_shifts

def load_schedule(schedule, month, year):
  print(f'year start: {year}')
  user_month = month
  user_year = year
  wordDoc = Document(schedule)
  # user_tz = pytz.timezone('America/Los_Angeles')
  while not user_month in month_variables:
    # user_month = simpledialog.askstring(title="Month", prompt="Please enter the 2-digit month")
    user_month = "08"

  while not int(user_year) > 2021:
    user_year = "2022"
  month = user_month
  # year = "2022"
  shifts = []
  shift_times = {2: "07:00", 3: "10:00", 4: "14:00", 5: "18:00"}
  for table in wordDoc.tables:
      date = []
      j = 0
      for row in table.rows:
        row_text = ''
        i = 0
        k = 0
        shift = ''
        time = ''
        for cell in row.cells:
          row_text = row_text + cell.text + ","
          if cell.text in month_list:
            month = cell.text
          if cell.text.lower() in day_list:
            j = 0
            # print('reading cells')
          else:
            if j % 6 == 1:
                date.append(cell.text)
                # print(f'date {date}')
            # else:
            #   if i % 2 == 0:
            elif i % 2 == 0:
                shift = cell.text
                # if j % 6 == 2:
                #   time = "07:00"
                # if j % 6 == 3:
                #   time = "10:00"
                # if j % 6 == 4:
                #   time = "14:00"
                # if j % 6 == 5:
                #   time = "18:00"
              # print(user)
            elif cell.text != "":
              cell_user = cell.text
              try:
                shift_start = datetime.datetime(
                  int(user_year),
                  int(user_month),
                  int(date[i]),
                  int(shift_times[j % 6][:2]),
                  00,
                  00
                )
                shift_start = timezone.make_aware(shift_start, timezone=TIMEZONE)
                # print(shift_start)
                # print(f'user: {cell_user}')
                shifts.append({
                  "user": cell_user,
                  "shift": shift, 
                  "date": user_year + "-" + user_month + "-" + date[i], 
                  "time": shift_times[j % 6], 
                  "month": month,
                  "year": user_year,
                  "shift_start": str(shift_start),
                  "shift_end": str(shift_start + timedelta(hours=12)),
                })
              except ValueError:
                print(f"Invalid date: {user_year}-{user_month}-{date[i]}")
          i += 1

        # print(row_text, "row", j)
        # print("date row", date)
        j += 1
        i = 0
        k = 0
        if j % 6 == 0: 
          date = []

  # header = ['Subject', 'Start date', 'Start time']
  # with open(f'schedule_{user_month}-{user_year}.csv', 'w', encoding='UTF8', newline='\n') as f:
  #   writer = csv.writer(f)
  #   writer.writerow(header)
  #   for shift in shifts:
  #     writer.writerow(shift)

  # f.close()
  # start_time = datetime.datetime(2023,6,21,3,45,00)
  # end_time = start_time + timedelta(hours=12)
  # print("start: ", start_time, "end: ", end_time)
  # print(shifts)
  load_database(shifts, user_month, user_year)
  # events = add_shifts(shifts)
  # json_shifts = json.dumps(events)
  return "Success!"

def load_database(shifts, month, year):
  try:
    print(month, year)
    Calendar.objects.filter(month=month, year=year).delete()
    i = 0
    for shift in shifts:
      calendar_shift = Calendar.objects.create(
        user_initials = shift["user"].strip(),
        start = shift["shift_start"],
        end = shift["shift_end"],
        month = month,
        year = year,
      )
    return
  # except HttpError as error:
  #   print('An error occurred: %s' % error)  
  except Exception as e:
    return trace_error(e, True)


SCOPES = ['https://www.googleapis.com/auth/calendar']

def get_users(schedule):
  wordDoc = Document(schedule) # Identify schedule as a Word Doc
  user_list = []
  for table in wordDoc.tables:
    for row in table.rows:
      row_text = ''
      for cell in row.cells:
        row_text = row_text + cell.text + ","
        if cell.text.isalpha() and len(cell.text) == 2:
          # print(f'cell: {cell.text}')
          if cell.text not in user_list:          
            user_list.append(cell.text)
    # print(user_list)
    # return json.dumps(user_list)
    return user_list


def test_calendar():
    print('running test_calendar')
    """Shows basic usage of the Google Calendar API.
    Prints the start and name of the next 10 events on the user's calendar.
    """
    creds = None
    # The file gmail.json stores the user's access and refresh tokens, and is
    # created automatically when the authorization flow completes for the first
    # time.
    if os.path.exists(f'backend/tokens/{gmail}.json'):
        print('Loading credentials from file')
        creds = Credentials.from_authorized_user_file(f'backend/tokens/{gmail}.json', SCOPES)
    # If there are no (valid) credentials available, let the user log in.
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file(
                'backend/credentials.json', SCOPES)
            creds = flow.run_local_server(port=0)
        # Save the credentials for the next run
        with open(f'backend/tokens/{gmail}.json', 'w') as token:
            token.write(creds.to_json())

    try:
        service = build('calendar', 'v3', credentials=creds)

        # Call the Calendar API
        now = datetime.datetime.utcnow().isoformat() + 'Z'  # 'Z' indicates UTC time
        print('Getting the upcoming 10 events')
        events_result = service.events().list(calendarId='primary', timeMin=now,
                                              maxResults=10, singleEvents=True,
                                              orderBy='startTime').execute()
        events = events_result.get('items', [])

        if not events:
            print('No upcoming events found.')
            return

        # Prints the start and name of the next 10 events
        for event in events:
            start = event['start'].get('dateTime', event['start'].get('date'))
            print(start, event['summary'])
        return events

    except HttpError as error:
        print('An error occurred: %s' % error)

def test_event():
    print('running test_event')
    """Shows basic usage of the Google Calendar API.
    Prints the start and name of the next 10 events on the user's calendar.
    """
    creds = None
    # The file gmail.json stores the user's access and refresh tokens, and is
    # created automatically when the authorization flow completes for the first
    # time.
    if os.path.exists(f'backend/tokens/{gmail}.json'):
        print('Loading credentials from file')
        creds = Credentials.from_authorized_user_file(f'backend/tokens/{gmail}.json', SCOPES)
    # If there are no (valid) credentials available, let the user log in.
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file(
                'backend/credentials.json', SCOPES)
            creds = flow.run_local_server(port=0)
        # Save the credentials for the next run
        with open(f'backend/tokens/{gmail}.json', 'w') as token:
            token.write(creds.to_json())

    try:
        service = build('calendar', 'v3', credentials=creds)

        # Call the Calendar API
        now = datetime.datetime.utcnow().isoformat() + 'Z'  # 'Z' indicates UTC time
        event = {
          "summary": "Day Shift on June 21 7am-7pm",
          "location": "World Wide Web",
          "description": "Testing",
          "start": {
            "dateTime": "2023-06-21T07:00:00-07:00",
            "timeZone": "America/Los_Angeles"
          },
          "end": {
            "dateTime": "2023-06-21T19:00:00-07:00",
            "timeZone": "America/Los_Angeles"
          },
        }
        event = service.events().insert(calendarId='primary', body=event).execute()
        new_event = 'Event created: %s' % (event.get('htmlLink'))
        print(new_event)
        return new_event

    except HttpError as error:
        print('An error occurred: %s' % error)

def add_shifts(shifts):
  print('running add_events')
  creds = None
  # if os.path.exists(f'backend/tokens/{gmail}.json'):
  #     print('Loading credentials from file')
  #     creds = Credentials.from_authorized_user_file(f'backend/tokens/{gmail}.json', SCOPES)
  # # If there are no (valid) credentials available, let the user log in.
  # if not creds or not creds.valid:
  #     if creds and creds.expired and creds.refresh_token:
  #         creds.refresh(Request())
  #     else:
  #         flow = InstalledAppFlow.from_client_secrets_file(
  #             'backend/credentials.json', SCOPES)
  #         creds = flow.run_local_server(port=0)
  #     # Save the credentials for the next run
  #     with open(f'backend/tokens/{gmail}.json', 'w') as token:
  #         token.write(creds.to_json())

  try:
      # service = build('calendar', 'v3', credentials=creds)
      i = 0
      events = {}
      for shift in shifts:
        # print("first half: ", shift["shift_start"][:10])
        # print("second half: ", shift["shift_start"][11:])
        # events = {
        events[f'shift_{i}'] = {
          "summary": shift["user"], # shift["shift"] + "-" + shift["user"],
          "location": "AMCS",
          "description": shift["shift"],
          "start": {
            "dateTime": shift["shift_start"][:10] + "T" + shift["shift_start"][11:] + "-07:00",
            "timeZone": "America/Los_Angeles"
          },
          "end": {
            "dateTime": shift["shift_end"][:10] + "T" + shift["shift_end"][11:] + "-07:00",
            "timeZone": "America/Los_Angeles"
          },
        }
        i += 1
        # print(event["start"])
        # events = service.events().insert(calendarId='primary', body=event).execute()
      # print("new events: ", events)
      return events
  except HttpError as error:
    print('An error occurred: %s' % error)